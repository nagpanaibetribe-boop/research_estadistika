import os
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from scipy import stats
from docx import Document

from django.shortcuts import render
from django.http import HttpResponse
from django.conf import settings

from decouple import config
from openai import OpenAI

from .forms import UploadForm

# ================= OPENAI CLIENT =================
client = OpenAI(api_key=config("OPENAI_API_KEY"))

# ================= AI DISCUSSION =================
def ai_generate_discussion(title, description, table_data):
    try:
        response = client.responses.create(
            model="gpt-4.1-mini",
            input=f"""
You are an academic researcher writing thesis-style statistical interpretations.

{title}
{description}

Table Data:
{table_data}

Write a clear academic discussion (3–5 sentences).
"""
        )
        return response.output_text.strip()

    except Exception as e:
        print("OPENAI ERROR:", e)
        return "AI-generated discussion is unavailable at the moment."

# ================= UPLOAD PAGE =================
def index(request):
    return render(request, "index.html", {"form": UploadForm()})

# ================= ANALYZE =================
def analyze(request):
    if request.method == "POST":
        form = UploadForm(request.POST, request.FILES)
        if not form.is_valid():
            return render(request, "index.html", {"form": form})

        df = pd.read_excel(request.FILES["file"])
        total_n = len(df)

        raw_cols = df.columns.tolist()
        raw_data = df.values.tolist()

        numeric = df.select_dtypes(include=np.number)
        categorical = df.select_dtypes(exclude=np.number)

        table_discussions = []

        # ---------- RAW DATA ----------
        table_discussions.append((
            "Table A. Raw Data",
            ai_generate_discussion(
                "Table A. Raw Data",
                "This table presents the raw responses of all respondents.",
                raw_data
            )
        ))

        # ---------- STATISTICS ----------
        stats_table = []
        if not numeric.empty:
            stats_df = numeric.agg(["mean", "std"]).T.round(2)
            for v, r in stats_df.iterrows():
                stats_table.append([v, r["mean"], r["std"]])

            table_discussions.append((
                "Table B. Statistical Summary",
                ai_generate_discussion(
                    "Table B. Statistical Summary",
                    "This table summarizes central tendency and variability.",
                    stats_table
                )
            ))

        # ---------- FREQUENCY ----------
        freq_tables = []
        for col in categorical.columns:
            freq = df[col].value_counts(dropna=False).reset_index()
            freq.columns = ["Category", "Frequency"]
            freq["Percent"] = round(freq["Frequency"] / total_n * 100, 2)
            rows = freq.values.tolist()
            freq_tables.append((col, rows))

            table_discussions.append((
                f"Table C. Frequency Distribution of {col}",
                ai_generate_discussion(
                    f"Table C. Frequency Distribution of {col}",
                    "This table presents the distribution of categorical responses.",
                    rows
                )
            ))

        # ---------- CORRELATION ----------
        corr_cols, corr_rows = None, None
        if numeric.shape[1] > 1:
            corr = numeric.corr().round(3)
            corr_cols = corr.columns.tolist()
            corr_rows = [(i, r.tolist()) for i, r in corr.iterrows()]

            table_discussions.append((
                "Correlation Matrix",
                ai_generate_discussion(
                    "Correlation Matrix",
                    "This matrix shows relationships between numeric variables.",
                    corr_rows
                )
            ))

        # ---------- T-TEST ----------
        ttest_results = []
        cols = numeric.columns.tolist()
        for i in range(len(cols)):
            for j in range(i + 1, len(cols)):
                t, p = stats.ttest_ind(
                    numeric[cols[i]].dropna(),
                    numeric[cols[j]].dropna(),
                    equal_var=False
                )
                ttest_results.append({
                    "var1": cols[i],
                    "var2": cols[j],
                    "t_stat": round(t, 4),
                    "p_val": round(p, 4)
                })

        if ttest_results:
            table_discussions.append((
                "T-Test Results",
                ai_generate_discussion(
                    "T-Test Results",
                    "This table presents independent sample t-test results.",
                    ttest_results
                )
            ))

        # ---------- SAVE SESSION ----------
        request.session["report_data"] = {
            "title": form.cleaned_data["title"],
            "objective": form.cleaned_data["objective"],
            "problem": form.cleaned_data["problem"],
            "total_n": total_n,
            "raw_cols": raw_cols,
            "raw_data": raw_data,
            "stats_table": stats_table,
            "freq_tables": freq_tables,
            "corr_cols": corr_cols,
            "corr_rows": corr_rows,
            "ttest_results": ttest_results,
            "table_discussions": table_discussions,
        }

        return render(request, "results.html", request.session["report_data"])

# ================= EXPORT WORD =================
def export_word(request):
    data = request.session.get("report_data")
    if not data:
        return HttpResponse("No data available.")

    doc = Document()
    doc.add_heading(data["title"], level=1)
    doc.add_paragraph(f"Objective: {data['objective']}")
    doc.add_paragraph(f"Problem: {data['problem']}")
    doc.add_paragraph(f"Respondents: {data['total_n']}")

    doc.add_heading("Discussion of Results", level=1)
    for title, discussion in data["table_discussions"]:
        doc.add_heading(title, level=2)
        doc.add_paragraph(discussion)

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = "attachment; filename=Research_Report.docx"
    doc.save(response)
    return response
